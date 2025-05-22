import sys
import PySimpleGUI as sg
import subprocess
import os
import os.path
from docx import Document
from docx.shared import Inches

# uncomment the following 3 lines before creating as an executable, this code hides the cmd window when run as .exe
#import win32gui, win32con
#The_program_to_hide = win32gui.GetForegroundWindow()
#win32gui.ShowWindow(The_program_to_hide , win32con.SW_HIDE)

def Launcher():
    def print(line):
            window['output'].update(line)
    sg.theme('black')
    
    # ------ Menu Definition ------ #
    menu_def = [['&Settings', ['&Team Paths', ['&Update Avengers', '&Update BOD', '&Update Gatekeepers','&Update Fellowship', '&Update IAM Legend', '&Update Nerd Herd', '&Update SuperCoP', '&Update Custom'], 
                ['docx Options', ['Show', 'Hide']], 'E&xit', 'Properties']],
    ]

    sg.set_options(element_padding=(1, 1), button_element_size=(10, 1), auto_size_buttons=False)

    col1 = [[sg.Text('Save Options', font='Calibri 8')], 
                [sg.Text(' ' * 4), sg.Text('Filename: ', text_color='white'), sg.Input(size=(20,1), key='-vFilename-')],
                [sg.Text(' ' * 2), sg.Text('User Story: ', text_color='white'), sg.Input(size=(20,1), key=('-vFolder2-'))],
                [sg.Text(' ' * 9), sg.Text('Team: ', text_color='white'), sg.OptionMenu(('Avengers', 'BOD', 'Gatekeepers', 'Fellowship', 'IAM Legend', 'Nerd Herd', 'SuperCoP', 'Desktop', 'Custom'),default_value='IAM Legend', size=(15,1), key='-vPath-'),sg.Text('   '),],
                [sg.Text(' ' * 8), sg.Text('Monitor'), sg.OptionMenu(('1', '2'), key='-vDisplay-'), sg.Text(' ' * 6),
                sg.Checkbox('Delay',key='-delay-')],
                #[sg.Input(size=(12,1), key='-vFolder-'), sg.FolderBrowse(button_color=('white','black'), font='Calibri 10', size=(7,1)),],
                [sg.Text(' ', size=(30,1), font='Calibri 8', text_color='red', key='output')],
    ]
    col2 =   [[sg.Text(' ' * 1), sg.Button(image_filename='./files/cam5_100x97.png', image_size=(100,97), border_width='0', key='snip'),],
    ]
    col3 =  [[sg.Checkbox('Add to docx', text_color='white', key='-appenddocx-'), sg.Text(' ' * 5), 
                sg.Input(size=(20,1), key='-vDocx-'), sg.FileBrowse(button_color=('white','black'), font='Calibri 10', size=(7,1)), 
                sg.Text(' ' * 5), #sg.Text('Monitor', text_color='white',), sg.Spin(values=('1', '2', '3'), size=(3,1), initial_value='1', key='-vDisplay-'), ],
                ],
    ]
    layout =  [
        [sg.Menu(menu_def, tearoff=True)],
        [sg.Text('QuickSnip', text_color='white', font='Forte 14'), sg.Text(' ' * 30),
        sg.Button('Clear',button_color=('white', '#000000'), size=(5,1), key='-clear-'),
        sg.Button('+', size=(4,1), button_color=('white', '#000000')), sg.Button('--', size=(4,1), button_color=('white', '#000000')), sg.Button('x', size=(4,1), button_color=('white', '#fc0303'))],
        [sg.Frame(layout= [                
        [sg.Column(col2, key='col2', visible=True), sg.Column(col1, key='col1', visible=True)],
        [sg.Column(col3, key='col3', visible=True)],], title='', border_width=(0), key='frame')]
    ]

    window = sg.Window('QuickSnip', layout, size=(400, 180), no_titlebar='none', grab_anywhere=True, keep_on_top=True, resizable=False)

    # ---===--- Loop taking in user input and executing appropriate program ---===--- #
    while True:
        event, values = window.read()
        if event == 'x' or event == 'Exit' or event is None:
            break           # exit button clicked
    # ---===--- GUI Window and Controls ---===--- #
        if event == '--':                                   #Minimize window
            window.Element('frame').Update(visible=False)
            window.Refresh()
            window.Size = (400,35)
        if event == '+':                                    #Maximize window
            window.Element('frame').Update(visible=True)
            window.Refresh()
            window.Refresh()
            if values['-appenddocx-'] == True:
                window.Size = (400,210)
            else:
                window.Size = (400, 180)
        if event == '-clear-':                              #Clear Window
            window['-vFilename-'].Update('')
            window['-vFolder2-'].Update('')
            window['-vDocx-'].Update('')
            window.Refresh()
        if event == 'Show':                                 #Show docx dropdown
            window['-appenddocx-'].Update(True)
            window.Element('frame').Update(visible=True)
            window.size = (400, 210)
            window.Refresh()
        if event == 'Hide':                                 #Hide docx dropdown
            window['-appenddocx-'].Update(False)
            window.Element('frame').Update(visible=True)
            window.size = (400, 180)
            window.Refresh()
    # ---===--- Team Save File Paths Update ---===--- #
        if event == 'Update Avengers':
            vAv_path = sg.popup_get_folder('Where would you like to save Avengers screenshots?', title='Browse for folder')
            if vAv_path is not None:
                chk_fldr = os.path.exists(vAv_path)    #Check Path, if folder doesn't exist create the folder
                if chk_fldr == True:                   #check if folder exists, if not then create the folder
                    with open('./config_files/p4.txt', 'w') as file:
                        file.write(str(vAv_path))
                        file.close
                    print('Avengers path was saved!')
                    window.Refresh()
                    window.read(timeout=5000)
                    print(' ')
                else:                                   #If folder doesn't exist, exit script
                    print('ERROR: Folder location does not exist')
                    window.Refresh()
                    window.read(timeout=3000)
                    print(' ')
            else:
                print('Cancelled...no changes were made')
                window.Refresh()
                window.read(timeout=3000)
                print(' ')
        if event == 'Update BOD':
            vBOD_path = sg.popup_get_folder('Where would you like to save BOD screenshots?', title='Browse for folder')
            if vBOD_path is not None:
                chk_fldr = os.path.exists(vBOD_path)    #Check Path, if folder doesn't exist create the folder
                if chk_fldr == True:                   #check if folder exists, if not then create the folder
                    with open('./config_files/p6.txt', 'w') as file:
                        file.write(str(vBOD_path))
                        file.close
                    print('BOD path was saved!')
                    window.Refresh()
                    window.read(timeout=5000)
                    print(' ')
                else:                                   #If folder doesn't exist, exit script
                    print('ERROR: Folder location does not exist')
                    window.Refresh()
                    window.read(timeout=3000)
                    print(' ')
            else:
                print('Cancelled...no changes were made')
                window.Refresh()
                window.read(timeout=3000)
                print(' ')
        if event == 'Update Custom':
            vCustom_path = sg.popup_get_folder('Where would you like to save Custom screenshots?', title='Browse for folder')
            if vCustom_path is not None:
                chk_fldr = os.path.exists(vCustom_path)    #Check Path, if folder doesn't exist create the folder
                if chk_fldr == True:                   #check if folder exists, if not then create the folder
                    with open('./config_files/p7.txt', 'w') as file:
                        file.write(str(vCustom_path))
                        file.close
                    print('Custom path was saved!')
                    window.Refresh()
                    window.read(timeout=5000)
                    print(' ')
                else:                                   #If folder doesn't exist, exit script
                    print('ERROR: Folder location does not exist')
                    window.Refresh()
                    window.read(timeout=3000)
                    print(' ')
            else:
                print('Cancelled...no changes were made')
                window.Refresh()
                window.read(timeout=3000)
                print(' ')
        if event == 'Update Gatekeepers':
            vDc_path = sg.popup_get_folder('Where would you like to save Gatekeepers screenshots?', title='Browse for folder')
            if vDc_path is not None:
                chk_fldr = os.path.exists(vDc_path)    #Check Path, if folder doesn't exist create the folder
                if chk_fldr == True:                   #check if folder exists, if not then create the folder
                    with open('./config_files/p5.txt', 'w') as file:
                        file.write(str(vDc_path))
                        file.close
                    print('Gatekeepers path was saved!')
                    window.Refresh()
                    window.read(timeout=5000)
                    print(' ')
                else:                                   #If folder doesn't exist, exit script
                    print('ERROR: Folder location does not exist')
                    window.Refresh()
                    window.read(timeout=3000)
                    print(' ')
            else:
                print('Cancelled...no changes were made')
                window.Refresh()
                window.read(timeout=3000)
                print(' ')
        if event == 'Update SuperCoP':
            vSc_path = sg.popup_get_folder('Where would you like to save SuperCoP screenshots?', title='Browse for folder')
            if vSc_path is not None:
                chk_fldr = os.path.exists(vSc_path)    #Check Path, if folder doesn't exist create the folder
                if chk_fldr == True:                   #check if folder exists, if not then create the folder
                    with open('./config_files/p1.txt', 'w') as file:
                        file.write(str(vSc_path))
                        file.close
                    print('SuperCop path was saved!')
                    window.Refresh()
                    window.read(timeout=5000)
                    print(' ')
                else:                                   #If folder doesn't exist, exit script
                    print('ERROR: Folder location does not exist')
                    window.Refresh()
                    window.read(timeout=3000)
                    print(' ')
            else:
                print('Cancelled...no changes were made')
                window.Refresh()
                window.read(timeout=3000)
                print(' ')
        if event == 'Update Nerd Herd':
            vNH_path = sg.popup_get_folder('Where would you like to save Nerd Herd screenshots?', title='Browse for folder')
            if vNH_path is not None:
                chk_fldr = os.path.exists(vNH_path)    #Check Path, if folder doesn't exist create the folder
                if chk_fldr == True:                   #check if folder exists, if not then create the folder
                    with open('./config_files/p2.txt', 'w') as file:
                        file.write(str(vNH_path))
                        file.close
                    print('Nerd Herd path was saved!')
                    window.Refresh()
                    window.read(timeout=5000)
                    print(' ')
                else:                                   #If folder doesn't exist, exit script
                    print('ERROR: Folder location does not exist')
                    window.Refresh()
                    window.read(timeout=3000)
                    print(' ')
            else:
                print('Cancelled...no changes were made')
                window.Refresh()
                window.read(timeout=3000)
                print(' ')
        if event == 'Update Fellowship':
            vFe_path = sg.popup_get_folder('Browse to Folder', 'Where would you like to save Fellowship screenshots?')
            if vFe_path is not None:
                chk_fldr = os.path.exists(vFe_path)    #Check Path, if folder doesn't exist create the folder
                if chk_fldr == True:                   #check if folder exists, if not then create the folder
                    with open('./config_files/p3.txt', 'w') as file:
                        file.write(str(vFe_path))
                        file.close
                    print('Fellowship path was saved!')
                    window.Refresh()
                    window.read(timeout=5000)
                    print(' ')
                else:                                   #If folder doesn't exist, exit script
                    print('ERROR: Folder location does not exist')
                    window.Refresh()
                    window.read(timeout=3000)
                    print(' ')
            else:
                print('Cancelled...no changes were made')
                window.Refresh()
                window.read(timeout=3000)
                print(' ')
        if event == 'Update IAM Legend':
            vIam_path = sg.popup_get_folder('Browse to Folder', 'Where would you like to save IAM Legend screenshots?')
            if vIam_path is not None:
                chk_fldr = os.path.exists(vIam_path)    #Check Path, if folder doesn't exist create the folder
                if chk_fldr == True:                   #check if folder exists, if not then create the folder
                    with open('./config_files/p8.txt', 'w') as file:
                        file.write(str(vIam_path))
                        file.close
                    print('IAM Legend path was saved!')
                    window.Refresh()
                    window.read(timeout=5000)
                    print(' ')
                else:                                   #If folder doesn't exist, exit script
                    print('ERROR: Folder location does not exist')
                    window.Refresh()
                    window.read(timeout=3000)
                    print(' ')
            else:
                print('Cancelled...no changes were made')
                window.Refresh()
                window.read(timeout=3000)
                print(' ')
    # ---===--- Screenshot Code ---===--- #
        if event == 'snip':
            try:
                if values['-delay-'] == True:
                    print('Starting in...')
                    window.read(timeout=1000)
                    print('3')
                    window.read(timeout=1000)
                    print('2')
                    window.read(timeout=1000)
                    print('1')
                    window.read(timeout=1000)
                print('Saving screenshot...please wait')
                window.Refresh()
                #####################################################################
                ############# Check which 'Team' is selected in the GUI #############
                ############# and set the path accordingly              #############
                #####################################################################
                if values['-vPath-'] == 'SuperCoP':     #Determine parent directory Path
                    f10 = open('./config_files/p1.txt', 'r')
                    p10 = f10.read()
                    vPath2 = p10
                    f10.close
                    #vPath2 = str(os.environ['USERPROFILE'] + vP3)
                elif values['-vPath-'] == 'Nerd Herd':
                    f20 = open('./config_files/p2.txt', 'r')
                    p20 = f20.read()
                    vPath2 = p20
                    f20.close
                    #vPath2 = str(os.environ['USERPROFILE'] + vP4)
                elif values['-vPath-'] == 'Fellowship':
                    f30 = open('./config_files/p3.txt', 'r')
                    p30 = f30.read()
                    vPath2 = p30
                    f30.close
                elif values['-vPath-'] == 'Avengers':
                    f40 = open('./config_files/p4.txt', 'r')
                    p40 = f40.read()
                    vPath2 = p40
                    f40.close
                elif values['-vPath-'] == 'DataCenter':
                    f50 = open('./config_files/p5.txt', 'r')
                    p50 = f50.read()
                    vPath2 = p50
                    f50.close
                    #vPath2 = str(os.environ['USERPROFILE'] + vP5)
                elif values['-vPath-'] == 'BOD':
                    f60 = open('./config_files/p6.txt', 'r')
                    p60 = f60.read()
                    vPath2 = p60
                    f60.close
                elif values['-vPath-'] == 'Custom':
                    f70 = open('./config_files/p7.txt', 'r')
                    p70 = f70.read()
                    vPath2 = p70
                    f70.close
                elif values['-vPath-'] == 'IAM Legend':
                    f80 = open('./config_files/p8.txt', 'r')
                    p80 = f80.read()
                    vPath2 = p80
                    f80.close
                elif values['-vPath-'] == 'Desktop':
                    vPath2 = str(r'U:\Desktop')
                #####################################################################
                ############# Insert folder between vPath & filname     #############
                #####################################################################
                vPath3 = values['-vFolder2-']
                if vPath3 is None:                      #If folder is blank, use parent directory path
                    vP = vPath2
                else:
                    vP = str(vPath2 + "\\" + vPath3)    #Combine str statements for complete folder path
                if vP.endswith("\\") == False:          #Path needs to end with \
                    vP = str(vP + "\\")                 
                #####################################################################
                ############# Check if folder exists, if not create it  #############
                #####################################################################
                chk_fldr = os.path.exists(vP)           #Check Path, if folder doesn't exist create the folder
                vF = values['-vFilename-']              #Screenshot name
                if chk_fldr == False:                   #check if folder exists, if not then create the folder
                        try:
                           os.mkdir(vP)
                        except:                         #If permission error occurs, end script
                            print('Permission Error: unable to create folder')
                            window.Refresh()
                #####################################################################
                #############     Determine which monitor to take a     #############
                #############     screenshot of, then call the script   #############
                #############     associated with that monitor          #############
                #####################################################################
                if values['-vDisplay-'] == '2': #Determine which monitor to take screenshot
                    with open("vF.txt", "w") as file:   #Create temp file to pass filename variable
                        file.write(str(vF))
                        file.close()
                    with open("vP.txt", "w") as file:   #Create temp file to pass partial path variable
                        file.write(str(vP))
                        file.close()
                    os.system('python QSSI2.py')        #Script to capture screenshot
                    os.remove("vF.txt")                 #Remove temp file
                    os.remove("vP.txt")                 #Remove temp file
                    print(' ')
                    window.refresh()
                elif values['-vDisplay-'] == '1':       #Determine which monitor to take screenshot
                    with open("vF.txt", "w") as file:   #Create temp file to pass full path variable
                        file.write(str(vP + vF))  
                        file.close()
                    os.system('python QSSI1.py')        #Script to capture screenshot
                    os.remove("vF.txt")                 #Remove temp file
                    print(' ')
                    window.refresh()
                #####################################################################
                #############  Add screenshot to docx if checkbox is    #############
                #############  checked. If docx file does not exist in  #############
                #############  target folder, then create & append it   #############
                #####################################################################
                if values['-appenddocx-'] == True:      #Append .docx file if box is checked
                    try:
                        print('Adding to Word doc...please wait')
                        window.refresh()
                        window.read(timeout=1500)
                        vDocx = str(vP + '\\' + values['-vDocx-'] + '.docx')
                        if os.path.isfile(vDocx):
                            print('Adding to Word doc...please wait')
                            window.refresh()
                            document = Document(vDocx)
                            document.save(vDocx)
                        else:
                            print('Creating Word doc')
                            window.refresh()
                            window.read(timeout=1500)
                            document = Document()
                            document.save(vDocx)
                        #####################################################################
                        #############   If screenshot file already exists, find #############
                        #############   last iteration of file                  #############
                        #####################################################################
                        vFilename2 = values['-vFilename-']                                    
                        picPath = str(vP + vFilename2 + '.png')
                        picPath2 = picPath.split(".png")[0] + "1.png"
                        if os.path.isfile(picPath): # file will be found since pic has already been taken
                            expand = 1
                            while True:
                                expand += 1
                                new_picPath = picPath.split(".png")[0] + str(expand) + ".png"
                                if os.path.isfile(new_picPath):
                                    continue
                                else:
                                    expand -= 1
                                    new_vF3 = picPath.split(".png")[0] + str(expand) + ".png"
                                    if new_vF3 == picPath2:    # if no iterator, then picPath remains unchanged
                                        picPath = str(vP + "\\" + vFilename2 + '.png')
                                        break
                                    else:
                                        picPath = new_vF3  # if iteretor used, last iterator will be picPath
                                        break
                        document = Document(vDocx)
                        p = document.add_paragraph(' ')
                        p = p.insert_paragraph_before('')
                        r = p.add_run()
                        #print(picPath)  
                        #window.refresh()
                        #window.read(timeout=3000)                          
                        r.add_picture(picPath, width=Inches(6))
                        p = p.insert_paragraph_before(vFilename2)
                        document.save(vDocx)
                        print(' ')
                    except:
                        print('ERROR: Screenshot saved but not added to docx')
                        window.refresh()
                        window.read(timeout=5000)
                        print(' ')
                elif values['-appenddocx-'] == False:
                        continue
                #break
            except:
                print('ERROR: Check filename and folder path')
                window.Refresh()
                window.read(timeout=5000)
                print(' ')

if __name__ == '__main__':
    Launcher()
